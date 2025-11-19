"""
Resume parsing service using LangChain and PyMuPDF
Extracts skills and experience level from resume files
"""

import os
import tempfile
import logging
from typing import Dict, List, Optional, Any
from pathlib import Path
import re

# Setup logger
logger = logging.getLogger(__name__)

# Try to import PyMuPDF, handle import errors and DLL errors gracefully
PYMUPDF_AVAILABLE = False
fitz = None
try:
    import fitz  # PyMuPDF
    # Test if it actually works (not just imported) - DLL might fail at runtime
    try:
        # Quick test: try to create an empty document
        test_doc = fitz.open()
        test_doc.close()
        PYMUPDF_AVAILABLE = True
    except Exception:
        # DLL error or other runtime error
        PYMUPDF_AVAILABLE = False
        fitz = None
except Exception as e:
    # Catch all exceptions including ImportError and DLL load errors on Windows
    PYMUPDF_AVAILABLE = False
    fitz = None

# Try to import python-docx, handle import errors gracefully
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    Document = None

# Note: LangChain components are not currently used in this parser
# They are kept for potential future use

class ResumeParser:
    """Parse resume files and extract skills and experience"""
    
    def __init__(self):
        self.skill_keywords = [
            # Programming Languages
            'python', 'java', 'javascript', 'typescript', 'c++', 'c#', 'go', 'rust', 'php', 'ruby',
            'swift', 'kotlin', 'scala', 'r', 'matlab', 'sql', 'html', 'css', 'sass', 'less',
            # Frameworks & Libraries
            'react', 'angular', 'vue', 'node.js', 'express', 'django', 'flask', 'fastapi',
            'spring', 'laravel', 'rails', 'asp.net', '.net', 'next.js', 'nuxt.js',
            # Databases
            'postgresql', 'mysql', 'mongodb', 'redis', 'elasticsearch', 'cassandra',
            'oracle', 'sqlite', 'dynamodb', 'firebase', 'supabase',
            
            # Cloud & DevOps
            'aws', 'azure', 'gcp', 'docker', 'kubernetes', 'jenkins', 'git', 'ci/cd',
            'terraform', 'ansible', 'linux', 'bash', 'shell scripting',
            # Tools & Others
            'git', 'github', 'gitlab', 'jira', 'agile', 'scrum', 'rest api', 'graphql',
            'microservices', 'machine learning', 'ai', 'data science', 'nlp', 'computer vision'
        ]
    
    def extract_text_from_pdf(self, file_path: str) -> str:
        """Extract text from PDF file using PyMuPDF with fallback"""
        if not os.path.exists(file_path):
            raise Exception(f"PDF file not found at path: {file_path}")
        
        # Check file size
        file_size = os.path.getsize(file_path)
        if file_size == 0:
            raise Exception("PDF file is empty (0 bytes)")
        
        print(f"[DEBUG] Attempting to parse PDF: {file_path} (size: {file_size} bytes)")
        
        # Try PyMuPDF first if available
        if PYMUPDF_AVAILABLE:
            try:
                # Open PDF in binary mode
                doc = fitz.open(file_path)
                print(f"[DEBUG] PDF opened successfully with PyMuPDF, pages: {len(doc)}")
                
                text = ""
                for page_num, page in enumerate(doc):
                    try:
                        page_text = page.get_text()
                        if page_text:
                            text += page_text + "\n"
                            print(f"[DEBUG] Extracted {len(page_text)} chars from page {page_num + 1}")
                    except Exception as page_error:
                        print(f"[WARNING] Error extracting text from page {page_num + 1}: {str(page_error)}")
                        continue
                
                doc.close()
                print(f"[DEBUG] Total text extracted: {len(text)} characters")
                
                if not text or len(text.strip()) < 10:
                    raise Exception("PDF file appears to be empty or contains no extractable text. The file might be image-based or corrupted.")
                
                return text
            except Exception as e:
                error_msg = str(e)
                print(f"[WARNING] PyMuPDF parsing failed: {error_msg}")
                print(f"[DEBUG] Falling back to PyPDF2...")
                # Continue to fallback below
        
        # Fallback: Try using PyPDF2 or pdfplumber if available
        try:
            import PyPDF2
            print("[DEBUG] Using PyPDF2 for PDF parsing")
            with open(file_path, 'rb') as pdf_file:
                pdf_reader = PyPDF2.PdfReader(pdf_file)
                print(f"[DEBUG] PDF opened with PyPDF2, pages: {len(pdf_reader.pages)}")
                text = ""
                for page_num, page in enumerate(pdf_reader.pages):
                    try:
                        page_text = page.extract_text() or ""
                        if page_text:
                            text += page_text + "\n"
                            print(f"[DEBUG] Extracted {len(page_text)} chars from page {page_num + 1} using PyPDF2")
                    except Exception as page_error:
                        print(f"[WARNING] Error extracting text from page {page_num + 1}: {str(page_error)}")
                        continue
                
                if not text or len(text.strip()) < 10:
                    raise Exception("PDF file appears to be empty or contains no extractable text.")
                
                print(f"[DEBUG] Total text extracted with PyPDF2: {len(text)} characters")
                return text
        except ImportError:
            print("[DEBUG] PyPDF2 not available, trying pdfplumber...")
            try:
                import pdfplumber
                print("[DEBUG] Using pdfplumber for PDF parsing")
                with pdfplumber.open(file_path) as pdf:
                    print(f"[DEBUG] PDF opened with pdfplumber, pages: {len(pdf.pages)}")
                    text = ""
                    for page_num, page in enumerate(pdf.pages):
                        try:
                            page_text = page.extract_text() or ""
                            if page_text:
                                text += page_text + "\n"
                                print(f"[DEBUG] Extracted {len(page_text)} chars from page {page_num + 1} using pdfplumber")
                        except Exception as page_error:
                            print(f"[WARNING] Error extracting text from page {page_num + 1}: {str(page_error)}")
                            continue
                        
                        if not text or len(text.strip()) < 10:
                            raise Exception("PDF file appears to be empty or contains no extractable text.")
                        
                        print(f"[DEBUG] Total text extracted with pdfplumber: {len(text)} characters")
                        return text
            except ImportError:
                raise Exception(
                    "PDF parsing libraries not available. Please install one of: "
                    "pip install pymupdf OR pip install PyPDF2 OR pip install pdfplumber"
                )
        except Exception as e:
            error_msg = str(e)
            error_msg_lower = error_msg.lower()  # Cache lowercased string
            print(f"[ERROR] Fallback PDF parsing error: {error_msg}")
            # Optimized: use cached lowercased string instead of multiple .lower() calls
            if "not a pdf" in error_msg_lower or "invalid" in error_msg_lower or "cannot read" in error_msg_lower:
                raise Exception("The file is not a valid PDF document. Please upload a valid PDF file.")
            raise Exception(f"Error extracting text from PDF: {str(e)}")
    
    def extract_text_from_docx(self, file_path: str) -> str:
        """Extract text from DOCX file"""
        if not os.path.exists(file_path):
            raise Exception(f"DOCX file not found at path: {file_path}")
        
        # Check file size
        file_size = os.path.getsize(file_path)
        if file_size == 0:
            raise Exception("DOCX file is empty (0 bytes)")
        
        print(f"[DEBUG] Attempting to parse DOCX: {file_path} (size: {file_size} bytes)")
        
        if not DOCX_AVAILABLE:
            raise Exception("python-docx is not available. Please install it with: pip install python-docx")
        
        try:
            doc = Document(file_path)
            print(f"[DEBUG] DOCX opened successfully")
            
            text_parts = []
            for paragraph in doc.paragraphs:
                if paragraph.text:
                    text_parts.append(paragraph.text)
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text:
                            text_parts.append(cell.text)
            
            text = "\n".join(text_parts)
            print(f"[DEBUG] Total text extracted: {len(text)} characters")
            
            if not text or len(text.strip()) < 10:
                raise Exception("DOCX file appears to be empty or contains no extractable text.")
            
            return text
        except Exception as e:
            error_msg = str(e)
            error_msg_lower = error_msg.lower()  # Cache lowercased string
            print(f"[ERROR] DOCX parsing error: {error_msg}")
            # Optimized: use cached lowercased string instead of multiple .lower() calls
            if "not a docx" in error_msg_lower or "invalid" in error_msg_lower or "corrupt" in error_msg_lower or "cannot open" in error_msg_lower:
                raise Exception("The file is not a valid DOCX document. Please upload a valid DOCX file.")
            raise Exception(f"Error extracting text from DOCX: {str(e)}")
    
    def extract_text(self, file_path: str, file_extension: str) -> str:
        """Extract text from resume file based on extension"""
        file_extension = file_extension.lower()
        
        if file_extension == '.pdf':
            return self.extract_text_from_pdf(file_path)
        elif file_extension in ['.docx', '.doc']:
            return self.extract_text_from_docx(file_path)
        else:
            raise ValueError(f"Unsupported file type: {file_extension}")
    
    def extract_skills(self, text: str) -> List[str]:
        """
        Extract skills from resume text
        Time Complexity: O(n*m) where n = text length, m = number of skills
        Space Complexity: O(k) where k = found skills (max 20)
        Optimization: Cache lowercased text, use set for O(1) membership check
        """
        text_lower = text.lower()  # Cache once
        found_skills = []
        found_skills_set = set()  # Use set for O(1) membership check
        
        for skill in self.skill_keywords:
            # Check for skill in various formats
            skill_lower = skill.lower()  # Cache skill lowercase
            # Optimized: text is already lowercased, no need for IGNORECASE flag
            pattern = rf'\b{re.escape(skill_lower)}\b'
            
            if re.search(pattern, text_lower):
                # Capitalize properly
                skill_formatted = skill.title() if '.' not in skill else skill.upper()
                if skill_formatted not in found_skills_set:
                    found_skills.append(skill_formatted)
                    found_skills_set.add(skill_formatted)
                    if len(found_skills) >= 20:  # Early exit when limit reached
                        break
        
        return found_skills[:20]  # Limit to top 20 skills
    
    def extract_experience_level(self, text: str) -> Optional[str]:
        """
        Extract experience level from resume text
        Time Complexity: O(n) where n = text length (single pass through patterns)
        Space Complexity: O(1)
        Optimization: Cache lowercased text, no IGNORECASE flag needed
        """
        text_lower = text.lower()  # Cache once
        
        # Patterns to match experience
        experience_patterns = [
            (r'\b(\d+)\s*(?:years?|yrs?|y\.?)\s*(?:of\s*)?experience', 'years'),
            (r'experience[:\s]+(\d+)\s*(?:years?|yrs?)', 'years'),
            (r'(\d+)\s*(?:years?|yrs?)\s*(?:in|of)', 'years'),
            (r'fresher|fresh\s*graduate|no\s*experience|entry\s*level', 'fresher'),
        ]
        
        # Check for fresher first (optimized: single pattern check)
        fresher_pattern = r'fresher|fresh\s*graduate|no\s*experience|entry\s*level'
        if re.search(fresher_pattern, text_lower):
            return "Fresher"
        
        # Check for years of experience
        # Optimized: compile patterns once and reuse
        max_years = 0
        for pattern, _ in experience_patterns:
            if pattern.startswith('fresher'):
                continue
            # Optimized: text is already lowercased, no IGNORECASE flag needed
            matches = re.finditer(pattern, text_lower)
            for match in matches:
                try:
                    years = int(match.group(1))
                    max_years = max(max_years, years)
                except (IndexError, ValueError):
                    continue
        
        if max_years > 0:
            return f"{max_years}yrs"
        
        # Check for keywords that might indicate experience level
        # Optimized: single regex with alternation instead of multiple searches
        if re.search(r'senior|lead|principal|architect', text_lower):
            return "5yrs+"
        elif re.search(r'mid|middle|intermediate', text_lower):
            return "2-4yrs"
        elif re.search(r'junior|entry|associate', text_lower):
            return "1yrs"
        
        return "Fresher"  # Default
    
    def extract_keywords(self, text: str) -> Dict[str, List[str]]:
        """Extract keywords including tools, technologies, and job titles"""
        text_lower = text.lower()
        keywords = {
            "tools": [],
            "technologies": [],
            "job_titles": [],
            "projects": []
        }
        
        # Extract tools and technologies (more comprehensive)
        tech_keywords = [
            # Frameworks
            'django', 'flask', 'fastapi', 'react', 'angular', 'vue', 'next.js', 'nuxt.js',
            'spring', 'express', 'laravel', 'rails', 'asp.net', '.net',
            # Tools
            'docker', 'kubernetes', 'jenkins', 'git', 'github', 'gitlab', 'jira', 'confluence',
            'terraform', 'ansible', 'puppet', 'chef', 'vagrant',
            # Databases
            'postgresql', 'mysql', 'mongodb', 'redis', 'elasticsearch', 'cassandra',
            'oracle', 'sqlite', 'dynamodb', 'firebase', 'supabase',
            # Cloud
            'aws', 'azure', 'gcp', 'heroku', 'vercel', 'netlify',
            # Other technologies
            'rest api', 'graphql', 'microservices', 'serverless', 'lambda',
            'machine learning', 'ai', 'data science', 'nlp', 'computer vision',
            'blockchain', 'web3', 'ethereum', 'solidity'
        ]
        
        for keyword in tech_keywords:
            pattern = rf'\b{re.escape(keyword.replace(".", "\\."))}\b'
            if re.search(pattern, text_lower, re.IGNORECASE):
                formatted = keyword.title() if '.' not in keyword else keyword
                if formatted not in keywords["technologies"]:
                    keywords["technologies"].append(formatted)
        
        # Extract job titles
        job_title_patterns = [
            r'(?:software|web|mobile|full.?stack|front.?end|back.?end|devops|data|ml|ai)\s+(?:engineer|developer|architect|specialist)',
            r'(?:senior|junior|lead|principal)\s+(?:software|web|mobile|full.?stack|front.?end|back.?end|devops|data|ml|ai)\s+(?:engineer|developer|architect)',
            r'(?:python|java|javascript|react|angular|vue|node)\s+(?:developer|engineer)',
            r'data\s+(?:scientist|engineer|analyst)',
            r'(?:machine\s+learning|ml|ai)\s+(?:engineer|scientist)',
            r'devops\s+(?:engineer|specialist)',
            r'system\s+(?:administrator|admin|architect)',
            r'qa\s+(?:engineer|tester|analyst)',
            r'product\s+(?:manager|owner)',
            r'technical\s+(?:lead|manager|architect)'
        ]
        
        for pattern in job_title_patterns:
            matches = re.finditer(pattern, text_lower, re.IGNORECASE)
            for match in matches:
                title = match.group(0).title()
                if title not in keywords["job_titles"]:
                    keywords["job_titles"].append(title)
        
        # Extract project mentions (simplified - looks for "project" followed by description)
        project_patterns = [
            r'project[:\s]+([A-Z][^.!?]{10,100})',
            r'built\s+([a-z\s]{10,80})\s+(?:using|with|in)',
            r'developed\s+([a-z\s]{10,80})\s+(?:using|with|in)',
            r'created\s+([a-z\s]{10,80})\s+(?:using|with|in)'
        ]
        
        for pattern in project_patterns:
            matches = re.finditer(pattern, text_lower, re.IGNORECASE)
            for match in matches:
                if len(match.groups()) > 0:
                    project = match.group(1).strip()[:80]
                    if project and project not in keywords["projects"]:
                        keywords["projects"].append(project)
        
        return keywords
    
    def extract_name(self, text: str) -> Optional[str]:
        """
        Extract name from resume text (usually at the beginning)
        Time Complexity: O(n) where n = number of lines checked (max 10)
        Space Complexity: O(1)
        Optimization: Cache excluded words set, optimize regex checks
        """
        lines = text.split('\n')[:10]  # Check first 10 lines
        # Optimized: cache excluded words in a set for O(1) lookup
        excluded_words = {'email', 'phone', 'address', 'resume', 'cv'}
        
        for line in lines:
            line = line.strip()
            if len(line) > 3 and len(line) < 50:  # Reasonable name length
                # Check if it looks like a name (contains letters, may have spaces)
                if re.match(r'^[A-Za-z\s\.\-]+$', line):
                    # Optimized: use set for O(1) membership check instead of list
                    line_words = {word.lower() for word in line.split()}
                    if not line_words.intersection(excluded_words):
                        # Check if it's not an email or phone
                        if '@' not in line and not re.match(r'^[\d\s\-\+\(\)]+$', line):
                            return line.title()
        return None
    
    def extract_email(self, text: str) -> Optional[str]:
        """Extract email address from resume text"""
        email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
        matches = re.findall(email_pattern, text)
        if matches:
            return matches[0].lower()
        return None
    
    def extract_projects(self, text: str) -> List[Dict[str, str]]:
        """Extract project information from resume text"""
        projects = []
        text_lower = text.lower()
        
        # Enhanced project extraction patterns
        project_patterns = [
            # Project title followed by description
            r'(?:project|project name|title)[:\s]+([A-Z][^.!?\n]{10,80})[:\s]*([^.!?\n]{20,200})',
            # Bullet points with project indicators
            r'(?:^|\n)\s*[-•*]\s*(?:project|built|developed|created|designed)[:\s]+([A-Z][^.!?\n]{10,100})',
            # Section headers like "PROJECTS" or "PROJECT EXPERIENCE"
            r'(?:projects?|project experience|personal projects?)[:\s]*\n((?:[^\n]{10,150}\n){1,5})',
        ]
        
        # Try to find project sections
        project_section_pattern = r'(?:projects?|project experience|personal projects?|notable projects?)[:\s]*\n((?:[^\n]+\n){2,10})'
        section_matches = re.finditer(project_section_pattern, text, re.IGNORECASE | re.MULTILINE)
        
        for section_match in section_matches:
            section_text = section_match.group(1)
            # Extract project names and descriptions from section
            lines = section_text.split('\n')
            current_project = None
            
            for line in lines:
                line = line.strip()
                if not line or len(line) < 10:
                    continue
                
                # Check if line looks like a project title (starts with capital, no bullet)
                if re.match(r'^[A-Z][^.!?]{5,60}(?:\s|$)', line) and not line.startswith(('-', '•', '*')):
                    if current_project:
                        projects.append(current_project)
                    current_project = {
                        "name": line[:80],
                        "summary": ""
                    }
                elif current_project and len(current_project["summary"]) < 200:
                    # Add to summary
                    if current_project["summary"]:
                        current_project["summary"] += " " + line[:150]
                    else:
                        current_project["summary"] = line[:150]
            
            if current_project:
                projects.append(current_project)
        
        # Fallback: extract from common patterns
        if not projects:
            fallback_patterns = [
                r'(?:built|developed|created|designed)\s+([A-Z][^.!?\n]{10,80})\s+(?:using|with|in|for)',
                r'([A-Z][A-Za-z\s]{5,40})\s+(?:project|application|system|platform)',
            ]
            
            for pattern in fallback_patterns:
                matches = re.finditer(pattern, text, re.IGNORECASE)
                for match in matches:
                    project_name = match.group(1).strip()[:80]
                    # Try to get context around the match
                    start = max(0, match.start() - 50)
                    end = min(len(text), match.end() + 150)
                    context = text[start:end]
                    summary = context.replace(project_name, "").strip()[:150]
                    
                    if project_name and len(project_name) > 5:
                        projects.append({
                            "name": project_name,
                            "summary": summary if summary else "Project details extracted from resume."
                        })
        
        # Remove duplicates and limit
        seen = set()
        unique_projects = []
        for proj in projects[:10]:  # Limit to 10 projects
            name_lower = proj["name"].lower()
            if name_lower not in seen:
                seen.add(name_lower)
                unique_projects.append(proj)
        
        return unique_projects
    
    def categorize_skills(self, skills: List[str], text: str) -> Dict[str, List[str]]:
        """Categorize skills into Programming, AI/ML, Tools, Soft Skills"""
        text_lower = text.lower()
        categorized = {
            "Programming": [],
            "AI/ML": [],
            "Tools": [],
            "Soft Skills": []
        }
        
        # Programming languages and frameworks
        programming_keywords = [
            'python', 'java', 'javascript', 'typescript', 'c++', 'c#', 'go', 'rust', 'php', 'ruby',
            'swift', 'kotlin', 'scala', 'r', 'matlab', 'sql', 'html', 'css', 'sass', 'less',
            'react', 'angular', 'vue', 'node.js', 'express', 'django', 'flask', 'fastapi',
            'spring', 'laravel', 'rails', 'asp.net', '.net', 'next.js', 'nuxt.js'
        ]
        
        # AI/ML keywords
        ai_ml_keywords = [
            'machine learning', 'ml', 'ai', 'artificial intelligence', 'deep learning',
            'neural network', 'tensorflow', 'pytorch', 'keras', 'scikit-learn',
            'nlp', 'natural language processing', 'computer vision', 'cv',
            'data science', 'data analysis', 'pandas', 'numpy', 'opencv'
        ]
        
        # Tools and DevOps
        tools_keywords = [
            'docker', 'kubernetes', 'jenkins', 'git', 'github', 'gitlab', 'jira',
            'terraform', 'ansible', 'linux', 'bash', 'shell', 'ci/cd',
            'postgresql', 'mysql', 'mongodb', 'redis', 'elasticsearch',
            'aws', 'azure', 'gcp', 'heroku', 'vercel', 'netlify'
        ]
        
        # Soft skills (extract from text, not from skill list typically)
        soft_skills_keywords = [
            'leadership', 'communication', 'teamwork', 'collaboration', 'problem solving',
            'agile', 'scrum', 'project management', 'mentoring', 'presentation'
        ]
        
        # Categorize each skill
        for skill in skills:
            skill_lower = skill.lower()
            categorized_flag = False
            
            # Check programming
            for keyword in programming_keywords:
                if keyword in skill_lower or skill_lower in keyword:
                    categorized["Programming"].append(skill)
                    categorized_flag = True
                    break
            
            if not categorized_flag:
                # Check AI/ML
                for keyword in ai_ml_keywords:
                    if keyword in skill_lower or skill_lower in keyword:
                        categorized["AI/ML"].append(skill)
                        categorized_flag = True
                        break
            
            if not categorized_flag:
                # Check tools
                for keyword in tools_keywords:
                    if keyword in skill_lower or skill_lower in keyword:
                        categorized["Tools"].append(skill)
                        categorized_flag = True
                        break
            
            # If not categorized, add to Programming as default
            if not categorized_flag:
                categorized["Programming"].append(skill)
        
        # Extract additional soft skills from text
        for keyword in soft_skills_keywords:
            if keyword in text_lower:
                formatted = keyword.title()
                if formatted not in categorized["Soft Skills"]:
                    categorized["Soft Skills"].append(formatted)
        
        # Remove duplicates within each category
        for category in categorized:
            categorized[category] = list(dict.fromkeys(categorized[category]))
        
        return categorized
    
    def generate_interview_topics(self, skills: List[str], keywords: Dict[str, List[str]], 
                                   projects: List[Dict[str, str]], text: str) -> List[str]:
        """Generate AI interview preparation topics based on resume content"""
        topics = set()
        text_lower = text.lower()
        
        # Core topics based on skills
        skill_to_topic = {
            'python': 'Python',
            'java': 'Java',
            'javascript': 'JavaScript',
            'typescript': 'TypeScript',
            'react': 'React',
            'angular': 'Angular',
            'vue': 'Vue.js',
            'node.js': 'Node.js',
            'django': 'Django',
            'flask': 'Flask',
            'fastapi': 'FastAPI',
            'sql': 'SQL',
            'postgresql': 'PostgreSQL',
            'mysql': 'MySQL',
            'mongodb': 'MongoDB',
            'redis': 'Redis',
            'docker': 'Docker',
            'kubernetes': 'Kubernetes',
            'aws': 'AWS',
            'azure': 'Azure',
            'gcp': 'GCP',
            'git': 'Git',
            'machine learning': 'Machine Learning',
            'ml': 'Machine Learning',
            'ai': 'Artificial Intelligence',
            'nlp': 'NLP',
            'data science': 'Data Science',
            'computer vision': 'Computer Vision',
            'tensorflow': 'TensorFlow',
            'pytorch': 'PyTorch'
        }
        
        # Add topics from skills
        for skill in skills:
            skill_lower = skill.lower()
            for key, topic in skill_to_topic.items():
                if key in skill_lower:
                    topics.add(topic)
        
        # Add topics from technologies
        for tech in keywords.get("technologies", []):
            tech_lower = tech.lower()
            for key, topic in skill_to_topic.items():
                if key in tech_lower:
                    topics.add(topic)
        
        # Add common interview topics based on domain
        if any(term in text_lower for term in ['machine learning', 'ml', 'ai', 'data science', 'nlp']):
            topics.update(['Machine Learning', 'Statistics', 'Linear Algebra', 'Probability'])
            if 'nlp' in text_lower or 'natural language' in text_lower:
                topics.add('NLP')
            if 'computer vision' in text_lower or 'cv' in text_lower:
                topics.add('Computer Vision')
        
        if any(term in text_lower for term in ['web', 'frontend', 'backend', 'full stack']):
            topics.update(['System Design', 'REST APIs', 'HTTP/HTTPS', 'Web Security'])
        
        if any(term in text_lower for term in ['database', 'sql', 'postgresql', 'mysql', 'mongodb']):
            topics.update(['Database Design', 'SQL Queries', 'Indexing', 'Transactions'])
        
        if any(term in text_lower for term in ['cloud', 'aws', 'azure', 'gcp']):
            topics.add('Cloud Architecture')
        
        if any(term in text_lower for term in ['docker', 'kubernetes', 'devops']):
            topics.update(['DevOps', 'CI/CD', 'Containerization'])
        
        # Always add fundamental topics
        topics.update(['Data Structures', 'Algorithms', 'Problem Solving'])
        
        # Add Git if mentioned
        if 'git' in text_lower:
            topics.add('Git')
        
        # Add APIs if mentioned
        if any(term in text_lower for term in ['api', 'rest', 'graphql']):
            topics.add('APIs')
        
        return sorted(list(topics))[:20]  # Limit to top 20 topics
    
    def calculate_resume_rating(self, name: Optional[str], email: Optional[str], 
                                skills: List[str], projects: List[Dict[str, str]],
                                experience_level: Optional[str], text_length: int,
                                keywords: Dict[str, List[str]]) -> float:
        """Calculate resume rating out of 5 based on multiple factors"""
        score = 0.0
        max_score = 5.0
        
        # Clarity (0.5 points) - name and email present
        if name:
            score += 0.25
        if email:
            score += 0.25
        
        # Skills (1.5 points) - based on number and diversity
        skills_score = min(len(skills) / 15.0, 1.0) * 1.5  # Max 1.5 for 15+ skills
        score += skills_score
        
        # Projects (1.0 point) - based on number and detail
        projects_score = min(len(projects) / 5.0, 1.0) * 1.0  # Max 1.0 for 5+ projects
        score += projects_score
        
        # Formatting/Completeness (1.0 point) - based on text length and structure
        if text_length > 1000:
            score += 0.5
        if text_length > 2000:
            score += 0.3
        if len(keywords.get("job_titles", [])) > 0:
            score += 0.2
        
        # Experience level (1.0 point) - if specified
        if experience_level and experience_level not in ["Not specified", "Unknown", "Fresher"]:
            score += 0.5
        elif experience_level == "Fresher":
            score += 0.2
        
        # Ensure score is between 0 and 5
        return round(min(max(score, 0.0), max_score), 2)
    
    def generate_resume_summary(self, name: Optional[str], email: Optional[str],
                               skills: List[str], experience_level: Optional[str],
                               keywords: Dict[str, List[str]], text: str) -> str:
        """Generate a clear resume summary paragraph"""
        summary_parts = []
        
        # Start with name if available
        if name:
            summary_parts.append(f"{name} is")
        else:
            summary_parts.append("This candidate is")
        
        # Add experience level
        if experience_level and experience_level not in ["Not specified", "Unknown"]:
            if experience_level == "Fresher":
                summary_parts.append("a fresher")
            else:
                summary_parts.append(f"an experienced professional with {experience_level} of experience")
        else:
            summary_parts.append("a professional")
        
        # Add domain/role
        job_titles = keywords.get("job_titles", [])
        if job_titles:
            primary_role = job_titles[0]
            summary_parts.append(f"specializing in {primary_role.lower()}")
        elif any(term in text.lower() for term in ['machine learning', 'ml', 'ai', 'data science']):
            summary_parts.append("with expertise in AI/ML and Data Science")
        elif any(term in text.lower() for term in ['web', 'frontend', 'backend']):
            summary_parts.append("with expertise in web development")
        elif any(term in text.lower() for term in ['mobile', 'android', 'ios']):
            summary_parts.append("with expertise in mobile development")
        else:
            summary_parts.append("with technical expertise")
        
        # Add key strengths
        strengths = []
        if len(skills) >= 10:
            strengths.append("strong technical skills")
        if len(keywords.get("technologies", [])) >= 5:
            strengths.append("diverse technology experience")
        if keywords.get("job_titles"):
            strengths.append("relevant industry experience")
        
        if strengths:
            summary_parts.append(f"demonstrating {', '.join(strengths)}")
        
        # Add technologies if available
        technologies = keywords.get("technologies", [])[:5]
        if technologies:
            tech_str = ", ".join(technologies[:3])
            summary_parts.append(f"with proficiency in {tech_str}")
        
        summary = " ".join(summary_parts) + "."
        return summary
    
    def generate_enhanced_summary(self, parsed_data: Dict[str, Any], text: str) -> Dict[str, Any]:
        """Generate comprehensive enhanced summary with all required components"""
        name = parsed_data.get("name")
        email = parsed_data.get("email")
        skills = parsed_data.get("skills", [])
        experience_level = parsed_data.get("experience_level")
        keywords = parsed_data.get("keywords", {})
        text_length = parsed_data.get("text_length", 0)
        
        # Extract projects
        projects = self.extract_projects(text)
        
        # Categorize skills
        categorized_skills = self.categorize_skills(skills, text)
        
        # Generate resume summary paragraph
        resume_summary = self.generate_resume_summary(name, email, skills, experience_level, keywords, text)
        
        # Generate interview topics
        interview_topics = self.generate_interview_topics(skills, keywords, projects, text)
        
        # Calculate rating
        rating = self.calculate_resume_rating(name, email, skills, projects, experience_level, text_length, keywords)
        
        return {
            "resume_summary": resume_summary,
            "skills_summary": categorized_skills,
            "projects_summary": projects,
            "interview_topics": interview_topics,
            "resume_rating": rating
        }
    
    def generate_interview_modules(self, parsed_data: Dict[str, Any], text: str) -> Dict[str, Any]:
        """Generate Interview Modules Overview with 4 modules - skills split based on resume"""
        skills = parsed_data.get("skills", [])
        experience_level = parsed_data.get("experience_level", "Fresher")
        keywords = parsed_data.get("keywords", {})
        # Get projects from summary if available, otherwise extract them
        summary = parsed_data.get("summary", {})
        projects = summary.get("projects_summary", []) if summary else []
        if not projects:
            projects = self.extract_projects(text)
        text_lower = text.lower()
        
        # Split skills into modules based on resume content
        technical_skills = self._extract_technical_skills(skills, keywords, text_lower)
        coding_topics = self._extract_coding_topics(skills, keywords, text_lower)
        hr_skills = self._extract_hr_skills(text_lower, keywords)
        star_points = self._extract_star_points(text, projects, keywords)
        
        # Technical Interview Module - core technical & domain skills
        technical_description = f"Focus on {len(technical_skills)} core technical skills and domain expertise extracted from your resume."
        
        # Coding / Online Test Module
        difficulty_level = self._determine_coding_difficulty(experience_level, skills, projects, text_lower)
        coding_platforms = self._recommend_coding_platforms(skills, text_lower)
        
        # HR Interview Module
        hr_evaluation_points = self._generate_hr_evaluation_points()
        
        # Behavioral Interview Module
        star_guidance = self._generate_star_guidance()
        
        return {
            "technical_interview": {
                "description": technical_description,
                "skills": technical_skills
            },
            "coding_test": {
                "difficulty_level": difficulty_level,
                "platforms": coding_platforms,
                "topics": coding_topics
            },
            "hr_interview": {
                "evaluation_points": hr_evaluation_points,
                "skills": hr_skills
            },
            "behavioral_interview": {
                "star_guidance": star_guidance,
                "star_points": star_points
            }
        }
    
    def _extract_technical_skills(self, skills: List[str], keywords: Dict[str, List[str]], text_lower: str) -> List[str]:
        """Extract core technical and domain skills for Technical Interview"""
        technical_skills = []
        
        # Core technical skills from extracted skills
        technical_keywords = [
            'python', 'java', 'javascript', 'typescript', 'react', 'angular', 'vue',
            'django', 'flask', 'fastapi', 'node.js', 'spring', 'sql', 'postgresql',
            'mysql', 'mongodb', 'redis', 'docker', 'kubernetes', 'aws', 'azure', 'gcp',
            'machine learning', 'ml', 'ai', 'data science', 'nlp', 'tensorflow', 'pytorch'
        ]
        
        for skill in skills:
            skill_lower = skill.lower()
            if any(keyword in skill_lower for keyword in technical_keywords):
                technical_skills.append(skill)
        
        # Add technologies from keywords
        technologies = keywords.get("technologies", [])
        for tech in technologies:
            if tech not in technical_skills:
                technical_skills.append(tech)
        
        # Add domain-specific skills based on text
        if any(term in text_lower for term in ['web', 'frontend', 'backend', 'full stack']):
            if 'Web Development' not in technical_skills:
                technical_skills.append('Web Development')
        
        if any(term in text_lower for term in ['machine learning', 'ml', 'ai', 'data science']):
            if 'AI/ML' not in technical_skills:
                technical_skills.append('AI/ML')
        
        if any(term in text_lower for term in ['mobile', 'android', 'ios', 'react native']):
            if 'Mobile Development' not in technical_skills:
                technical_skills.append('Mobile Development')
        
        return list(dict.fromkeys(technical_skills))[:20]  # Remove duplicates, limit to 20
    
    def _extract_coding_topics(self, skills: List[str], keywords: Dict[str, List[str]], text_lower: str) -> List[str]:
        """Extract programming, DSA, and problem-solving topics for Coding Test"""
        coding_topics = []
        
        # Programming languages
        programming_langs = ['python', 'java', 'javascript', 'c++', 'c#', 'go', 'rust']
        for skill in skills:
            skill_lower = skill.lower()
            if any(lang in skill_lower for lang in programming_langs):
                coding_topics.append(skill)
        
        # DSA topics based on skills and text
        if any(term in text_lower for term in ['algorithm', 'data structure', 'dsa', 'leetcode', 'hackerrank']):
            coding_topics.extend(['Data Structures', 'Algorithms', 'Problem Solving'])
        
        # Add common DSA topics if programming skills exist
        if coding_topics:
            coding_topics.extend(['Arrays & Strings', 'Hash Tables', 'Two Pointers', 'Dynamic Programming'])
        
        # Database-related coding
        if any(term in text_lower for term in ['sql', 'database', 'query']):
            coding_topics.append('SQL Queries')
        
        # System design if mentioned
        if any(term in text_lower for term in ['system design', 'architecture', 'scalability']):
            coding_topics.extend(['System Design', 'Scalability'])
        
        return list(dict.fromkeys(coding_topics))[:15]  # Remove duplicates, limit to 15
    
    def _extract_hr_skills(self, text_lower: str, keywords: Dict[str, List[str]]) -> List[str]:
        """Extract communication, teamwork, leadership skills for HR Interview"""
        hr_skills = []
        
        # Extract soft skills from text
        soft_skill_patterns = {
            'communication': ['communication', 'communicate', 'presentation', 'presenting'],
            'teamwork': ['team', 'collaboration', 'collaborate', 'teamwork', 'group'],
            'leadership': ['lead', 'leader', 'leadership', 'manage', 'management', 'mentor', 'supervise'],
            'problem solving': ['problem solving', 'problem-solve', 'analytical', 'analysis'],
            'adaptability': ['adapt', 'flexible', 'agile', 'scrum'],
            'time management': ['time management', 'deadline', 'prioritize']
        }
        
        for skill_name, patterns in soft_skill_patterns.items():
            if any(pattern in text_lower for pattern in patterns):
                hr_skills.append(skill_name.title())
        
        # Extract from job titles if available
        job_titles = keywords.get("job_titles", [])
        for title in job_titles:
            title_lower = title.lower()
            if 'lead' in title_lower or 'senior' in title_lower or 'manager' in title_lower:
                if 'Leadership' not in hr_skills:
                    hr_skills.append('Leadership')
        
        # Add strengths/weaknesses insights
        if len(hr_skills) >= 3:
            hr_skills.append('Strong interpersonal skills')
        if any(term in text_lower for term in ['project', 'deliver', 'achieve']):
            hr_skills.append('Project delivery experience')
        
        return list(dict.fromkeys(hr_skills))[:10]  # Remove duplicates, limit to 10
    
    def _extract_star_points(self, text: str, projects: List[Dict[str, str]], keywords: Dict[str, List[str]]) -> List[str]:
        """Extract STAR method points from resume for Behavioral Interview"""
        star_points = []
        
        # Extract from projects
        for project in projects:
            project_name = project.get("name", "")
            project_summary = project.get("summary", "")
            
            if project_name:
                star_points.append(f"Situation: Worked on {project_name}")
            
            if project_summary:
                # Try to extract action verbs
                action_verbs = ['developed', 'created', 'built', 'designed', 'implemented', 'optimized', 'improved']
                for verb in action_verbs:
                    if verb in project_summary.lower():
                        star_points.append(f"Action: {verb.capitalize()} solution for {project_name}")
                        break
        
        # Extract from text - look for achievement patterns
        achievement_patterns = [
            r'increased\s+[^.!?]{5,50}',
            r'reduced\s+[^.!?]{5,50}',
            r'improved\s+[^.!?]{5,50}',
            r'achieved\s+[^.!?]{5,50}',
            r'led\s+[^.!?]{5,50}',
            r'managed\s+[^.!?]{5,50}'
        ]
        
        for pattern in achievement_patterns:
            matches = re.finditer(pattern, text, re.IGNORECASE)
            for match in matches:
                result_text = match.group(0).strip()
                if len(result_text) < 100:  # Keep it concise
                    star_points.append(f"Result: {result_text}")
        
        # Extract challenges/situations
        challenge_keywords = ['challenge', 'problem', 'issue', 'difficult', 'complex']
        sentences = text.split('.')
        for sentence in sentences[:20]:  # Check first 20 sentences
            sentence_lower = sentence.lower()
            if any(keyword in sentence_lower for keyword in challenge_keywords) and len(sentence) > 20:
                star_points.append(f"Situation: {sentence.strip()[:80]}")
                break
        
        # Limit and ensure uniqueness
        return list(dict.fromkeys(star_points))[:8]  # Remove duplicates, limit to 8
    
    def _determine_coding_difficulty(self, experience_level: str, skills: List[str], 
                                    projects: List[Dict[str, str]], text_lower: str) -> str:
        """Determine coding test difficulty level"""
        if not experience_level or experience_level in ["Fresher", "Not specified", "Unknown"]:
            return "Easy to Medium"
        
        # Check for years of experience
        years_match = re.search(r'(\d+)', str(experience_level))
        years = int(years_match.group(1)) if years_match else 0
        
        if years >= 5:
            return "Hard"
        elif years >= 2:
            return "Medium to Hard"
        else:
            return "Easy to Medium"
    
    def _recommend_coding_platforms(self, skills: List[str], text_lower: str) -> List[str]:
        """Recommend coding platforms based on skills"""
        platforms = []
        
        # Always include common platforms
        platforms.extend(["LeetCode", "HackerRank", "CodeSignal"])
        
        # Add platform-specific recommendations
        if any(skill.lower() in ['python', 'java', 'javascript', 'c++'] for skill in skills):
            platforms.append("CodeChef")
        
        if any(term in text_lower for term in ['algorithm', 'data structure', 'competitive']):
            platforms.append("Codeforces")
        
        if any(term in text_lower for term in ['system design', 'architecture', 'distributed']):
            platforms.append("Pramp")
        
        return list(dict.fromkeys(platforms))[:5]  # Remove duplicates, limit to 5
    
    def _recommend_coding_topics(self, skills: List[str], keywords: Dict[str, List[str]], 
                                text_lower: str) -> List[str]:
        """Recommend coding topics for practice"""
        topics = []
        
        # Core topics
        topics.extend(["Arrays & Strings", "Hash Tables", "Two Pointers", "Sliding Window"])
        
        # Based on skills
        if any('python' in s.lower() or 'java' in s.lower() for s in skills):
            topics.extend(["Dynamic Programming", "Greedy Algorithms"])
        
        if any('javascript' in s.lower() or 'react' in s.lower() for s in skills):
            topics.extend(["Tree Traversal", "Graph Algorithms"])
        
        # Based on domain
        if any(term in text_lower for term in ['database', 'sql', 'postgresql', 'mysql']):
            topics.append("SQL Queries")
        
        if any(term in text_lower for term in ['system design', 'architecture', 'microservices']):
            topics.extend(["System Design", "Scalability"])
        
        if any(term in text_lower for term in ['machine learning', 'ml', 'ai', 'data science']):
            topics.extend(["Algorithm Optimization", "Time Complexity"])
        
        return list(dict.fromkeys(topics))[:8]  # Remove duplicates, limit to 8
    
    def _generate_hr_evaluation_points(self) -> List[str]:
        """Generate common HR evaluation points"""
        return [
            "Communication skills and clarity of expression",
            "Cultural fit and alignment with company values",
            "Motivation and interest in the role",
            "Career goals and long-term aspirations",
            "Salary expectations and negotiation readiness",
            "Availability and notice period",
            "Team collaboration and interpersonal skills"
        ]
    
    def _generate_hr_suggestions(self, experience_level: str, skills: List[str], 
                                projects: List[Dict[str, str]], keywords: Dict[str, List[str]],
                                text_lower: str) -> List[str]:
        """Generate personalized HR interview suggestions"""
        suggestions = []
        
        # Experience-based suggestions
        if experience_level and experience_level not in ["Fresher", "Not specified", "Unknown"]:
            suggestions.append(f"Highlight your {experience_level} of experience and key achievements.")
        else:
            suggestions.append("Emphasize your learning ability, projects, and academic achievements.")
        
        # Skills-based suggestions
        if len(skills) >= 10:
            suggestions.append("Showcase your diverse technical skill set and adaptability.")
        elif len(skills) >= 5:
            suggestions.append("Focus on your core competencies and depth of knowledge.")
        
        # Project-based suggestions
        if len(projects) >= 3:
            suggestions.append("Prepare to discuss your projects in detail, focusing on impact and learnings.")
        elif len(projects) > 0:
            suggestions.append("Be ready to explain your project contributions and problem-solving approach.")
        
        # Domain-specific suggestions
        if any(term in text_lower for term in ['leadership', 'lead', 'mentor', 'team']):
            suggestions.append("Highlight your leadership experience and team collaboration skills.")
        
        if any(term in text_lower for term in ['startup', 'entrepreneur', 'founder']):
            suggestions.append("Emphasize your entrepreneurial mindset and ability to work in fast-paced environments.")
        
        # Default suggestions
        if not suggestions:
            suggestions.append("Prepare clear examples of your work and be ready to discuss your career goals.")
            suggestions.append("Research the company culture and align your answers with their values.")
        
        return suggestions[:4]  # Limit to 4 suggestions
    
    def _generate_star_guidance(self) -> Dict[str, str]:
        """Generate STAR method guidance"""
        return {
            "Situation": "Set the context: Describe the situation or challenge you faced. Be specific about when and where this occurred.",
            "Task": "Explain your responsibility: What was your role? What needed to be accomplished?",
            "Action": "Detail your actions: What specific steps did you take? Focus on your contributions, not the team's.",
            "Result": "Share the outcome: What was the result? Quantify achievements when possible (e.g., 'increased efficiency by 30%')."
        }
    
    def _generate_behavioral_tips(self, experience_level: str, projects: List[Dict[str, str]],
                                  keywords: Dict[str, List[str]], text_lower: str) -> List[str]:
        """Generate personalized behavioral interview tips"""
        tips = []
        
        # Experience-based tips
        if experience_level and experience_level not in ["Fresher", "Not specified", "Unknown"]:
            tips.append("Prepare 3-5 detailed STAR stories from your professional experience.")
        else:
            tips.append("Use academic projects, internships, or personal projects as STAR examples.")
        
        # Project-based tips
        if len(projects) >= 2:
            tips.append("Select your most impactful projects and prepare detailed STAR narratives for each.")
        
        # Common behavioral questions preparation
        tips.append("Practice answering: 'Tell me about a time you faced a challenge' and 'Describe a conflict you resolved'.")
        
        # Leadership/teamwork tips
        if any(term in text_lower for term in ['team', 'collaboration', 'group project']):
            tips.append("Prepare examples demonstrating teamwork, conflict resolution, and collaboration.")
        
        if any(term in text_lower for term in ['lead', 'manage', 'mentor', 'supervise']):
            tips.append("Have ready examples of leadership, decision-making, and mentoring others.")
        
        # Problem-solving tips
        if any(term in text_lower for term in ['problem', 'solve', 'debug', 'fix']):
            tips.append("Prepare stories showing your analytical thinking and problem-solving process.")
        
        # Default tips
        if len(tips) < 4:
            tips.append("Quantify your achievements with numbers, percentages, or time saved.")
            tips.append("Be specific and avoid vague answers. Use concrete examples.")
        
        return tips[:5]  # Limit to 5 tips
    
    def parse_resume(self, file_path: str, file_extension: str) -> Dict[str, Any]:
        """Parse resume and extract all relevant information"""
        try:
            # Extract text
            text = self.extract_text(file_path, file_extension)
            
            if not text or len(text.strip()) < 50:
                raise ValueError("Resume file appears to be empty or invalid")
            
            # Extract personal information
            name = self.extract_name(text)
            email = self.extract_email(text)
            
            # Extract skills
            skills = self.extract_skills(text)
            
            # Extract experience level
            experience_level = self.extract_experience_level(text)
            
            # Extract additional keywords
            keywords = self.extract_keywords(text)
            
            # Build parsed data
            parsed_data = {
                "name": name,
                "email": email,
                "skills": skills,
                "experience_level": experience_level,
                "keywords": keywords,
                "text_length": len(text),
                "extracted_text_preview": text[:500]  # First 500 chars for debugging
            }
            
            # Generate enhanced summary
            try:
                enhanced_summary = self.generate_enhanced_summary(parsed_data, text)
                parsed_data["summary"] = enhanced_summary
            except Exception as summary_error:
                # If summary generation fails, continue without it
                logger.warning(f"Failed to generate enhanced summary: {str(summary_error)}")
                parsed_data["summary"] = None
            
            return parsed_data
        except Exception as e:
            raise Exception(f"Error parsing resume: {str(e)}")

# Create global instance
resume_parser = ResumeParser()

