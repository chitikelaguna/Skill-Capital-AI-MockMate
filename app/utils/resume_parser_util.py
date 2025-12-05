"""
Robust Resume Parser Utility
Handles PDF and DOCX parsing with multiple fallback libraries
"""

import os
import re
import logging
import platform
from typing import Dict, List, Optional, Any
from pathlib import Path

# Setup logger
logger = logging.getLogger(__name__)

# PDF Parsing Libraries
PYMUPDF_AVAILABLE = False
PDFPLUMBER_AVAILABLE = False

try:
    import fitz  # PyMuPDF
    try:
        test_doc = fitz.open()
        test_doc.close()
        PYMUPDF_AVAILABLE = True
    except Exception:
        PYMUPDF_AVAILABLE = False
except (ImportError, Exception):
    # Catch all exceptions including DLL load errors on Windows
    PYMUPDF_AVAILABLE = False
    pass

try:
    import pdfplumber
    PDFPLUMBER_AVAILABLE = True
except ImportError:
    pass

# DOCX Parsing
DOCX_AVAILABLE = False
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    pass

def is_text_meaningful(text: str, min_length: int = 20) -> bool:
    """
    Check if extracted text is meaningful (not just metadata/page numbers).
    LaTeX PDFs often extract only page numbers or metadata, which we want to treat as 'no text'.
    
    Returns True if text appears to be meaningful resume content, False otherwise.
    
    Made less strict to avoid false positives for valid PDFs:
    - Reduced min_length from 50 to 20 characters
    - Reduced alphanumeric ratio threshold from 30% to 20%
    - Reduced keyword requirement from 2 to 1
    """
    if not text or len(text.strip()) < min_length:
        return False
    
    # Check if text contains mostly numbers, whitespace, or special characters (likely metadata)
    text_stripped = text.strip()
    
    # Count alphanumeric characters (letters and numbers)
    alphanumeric_count = sum(1 for c in text_stripped if c.isalnum())
    total_chars = len(text_stripped)
    
    if total_chars == 0:
        return False
    
    # If less than 20% of text is alphanumeric, it's probably not meaningful (reduced from 30%)
    alphanumeric_ratio = alphanumeric_count / total_chars
    if alphanumeric_ratio < 0.2:
        return False
    
    # Check for common resume keywords (if present, text is likely meaningful)
    resume_keywords = [
        'experience', 'education', 'skills', 'project', 'work', 'job',
        'university', 'college', 'degree', 'email', 'phone', 'address',
        'python', 'java', 'javascript', 'react', 'developer', 'engineer',
        'software', 'programming', 'technology', 'certification', 'achievement',
        'resume', 'name', 'contact', 'summary', 'objective', 'profile'
    ]
    text_lower = text.lower()
    keyword_count = sum(1 for keyword in resume_keywords if keyword in text_lower)
    
    # If we have at least 1 resume keyword, text is likely meaningful (reduced from 2)
    if keyword_count >= 1:
        return True
    
    # If text is long enough and has reasonable alphanumeric ratio, consider it meaningful
    # Reduced threshold from 50% to 40% for alphanumeric ratio
    if len(text_stripped) >= min_length and alphanumeric_ratio >= 0.4:
        return True
    
    # If text has at least 30 characters and 30% alphanumeric, accept it (more lenient)
    if len(text_stripped) >= 30 and alphanumeric_ratio >= 0.3:
        return True
    
    return False


def parse_pdf(file_path: str, use_ocr_fallback: bool = False) -> Dict[str, Any]:
    """
    Parse PDF file with multiple fallback libraries
    Returns structured data with name, email, skills, experience
    
    Uses PyMuPDF and pdfplumber for text extraction.
    Note: OCR support has been removed. For LaTeX PDFs, please export as PDF/A.
    """
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"PDF file not found: {file_path}")
    
    file_size = os.path.getsize(file_path)
    logger.info(f"[RESUME_PARSER] Starting PDF parsing: {file_path} (size: {file_size} bytes)")
    
    # Check library availability
    available_parsers = []
    if PYMUPDF_AVAILABLE:
        available_parsers.append("PyMuPDF")
    if PDFPLUMBER_AVAILABLE:
        available_parsers.append("pdfplumber")
    
    logger.info(f"[RESUME_PARSER] Available PDF parsers: {', '.join(available_parsers) if available_parsers else 'NONE'}")
    
    if not available_parsers:
        raise ImportError("No PDF parsing libraries are installed. Please install dependencies: pip install -r requirements.txt")
    
    text = None
    parser_used = None
    last_error = None
    text_is_meaningful = False
    
    # Try PyMuPDF first (most reliable)
    if PYMUPDF_AVAILABLE:
        try:
            logger.debug(f"[RESUME_PARSER] Attempting PyMuPDF parsing...")
            doc = fitz.open(file_path)
            text_parts = []
            for page_num, page in enumerate(doc):
                try:
                    page_text = page.get_text()
                    if page_text:
                        text_parts.append(page_text)
                        logger.debug(f"[RESUME_PARSER] PyMuPDF extracted {len(page_text)} chars from page {page_num + 1}")
                except Exception as page_error:
                    logger.warning(f"[RESUME_PARSER] PyMuPDF failed on page {page_num + 1}: {str(page_error)}")
            doc.close()
            text = "\n".join(text_parts)
            
            # Check if text is meaningful (not just metadata)
            if text:
                text_is_meaningful = is_text_meaningful(text, min_length=20)
                if text_is_meaningful:
                    parser_used = "PyMuPDF"
                    logger.info(f"[RESUME_PARSER] PyMuPDF SUCCESS - Extracted {len(text)} meaningful characters")
                else:
                    logger.warning(f"[RESUME_PARSER] PyMuPDF extracted {len(text)} chars but text appears to be metadata only (not meaningful)")
                    # Don't reset text immediately - try other parsers first
                    # Only reset if all parsers fail
        except Exception as e:
            last_error = str(e)
            logger.error(f"[RESUME_PARSER] PyMuPDF failed: {str(e)}")
    
    # Fallback to pdfplumber
    if not text or not text_is_meaningful:
        if PDFPLUMBER_AVAILABLE:
            try:
                logger.debug(f"[RESUME_PARSER] Attempting pdfplumber parsing...")
                with pdfplumber.open(file_path) as pdf:
                    text_parts = []
                    for page_num, page in enumerate(pdf.pages):
                        try:
                            page_text = page.extract_text()
                            if page_text:
                                text_parts.append(page_text)
                                logger.debug(f"[RESUME_PARSER] pdfplumber extracted {len(page_text)} chars from page {page_num + 1}")
                        except Exception as page_error:
                            logger.warning(f"[RESUME_PARSER] pdfplumber failed on page {page_num + 1}: {str(page_error)}")
                    text = "\n".join(text_parts)
                
                if text:
                    text_is_meaningful = is_text_meaningful(text, min_length=20)
                    if text_is_meaningful:
                        parser_used = "pdfplumber"
                        logger.info(f"[RESUME_PARSER] pdfplumber SUCCESS - Extracted {len(text)} meaningful characters")
                    else:
                        logger.warning(f"[RESUME_PARSER] pdfplumber extracted {len(text)} chars but text appears to be metadata only")
                        # Don't reset text - continue to next parser
            except Exception as e:
                last_error = str(e)
                logger.error(f"[RESUME_PARSER] pdfplumber failed: {str(e)}")
    
        # Final check: if we have text but it wasn't marked as meaningful, 
        # check one more time with more lenient criteria before giving up
        if text and not text_is_meaningful:
            # Try with even more lenient criteria (10 chars minimum)
            if len(text.strip()) >= 10:
                # If we have at least 10 characters, accept it even if not "meaningful"
                # This prevents false positives for simple but valid PDFs
                logger.info(f"[RESUME_PARSER] Accepting text with {len(text)} chars (lenient mode)")
                text_is_meaningful = True
                if not parser_used:
                    parser_used = "PyMuPDF"  # Default to first parser used
        
        # If still no meaningful text, raise error
        if not text or not text_is_meaningful:
            error_details = []
            if not available_parsers:
                error_details.append("No parsing libraries available")
            else:
                error_details.append(f"All parsers failed (tried: {', '.join(available_parsers)})")
            if last_error:
                error_details.append(f"Last error: {last_error}")
            
            logger.warning(f"[RESUME_PARSER] FAILED - No meaningful text found. {'. '.join(error_details)}")
            
            # Provide error message for LaTeX/vector PDFs
            raise ValueError("PDF_NO_TEXT: Could not extract meaningful text from PDF. The file might be a LaTeX-generated PDF (vector-based), image-based (scanned), corrupted, or password-protected. For LaTeX PDFs, please export as PDF/A from Overleaf or use a PDF with selectable text.")
    
    logger.info(f"[RESUME_PARSER] PDF parsing completed successfully using {parser_used}")
    return extract_resume_data(text)


def parse_docx(file_path: str) -> Dict[str, Any]:
    """
    Parse DOCX file
    Returns structured data with name, email, skills, experience
    """
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"DOCX file not found: {file_path}")
    
    file_size = os.path.getsize(file_path)
    logger.info(f"[RESUME_PARSER] Starting DOCX parsing: {file_path} (size: {file_size} bytes)")
    
    if not DOCX_AVAILABLE:
        logger.error("[RESUME_PARSER] python-docx is not installed")
        raise ImportError("python-docx is not installed. Please install dependencies: pip install -r requirements.txt")
    
    try:
        logger.debug(f"[RESUME_PARSER] Attempting DOCX parsing with python-docx...")
        doc = Document(file_path)
        text_parts = []
        
        # Extract from paragraphs
        paragraph_count = 0
        for paragraph in doc.paragraphs:
            if paragraph.text:
                text_parts.append(paragraph.text)
                paragraph_count += 1
        logger.debug(f"[RESUME_PARSER] Extracted text from {paragraph_count} paragraphs")
        
        # Extract from tables
        table_count = 0
        for table in doc.tables:
            table_count += 1
            for row in table.rows:
                for cell in row.cells:
                    if cell.text:
                        text_parts.append(cell.text)
        logger.debug(f"[RESUME_PARSER] Extracted text from {table_count} tables")
        
        text = "\n".join(text_parts)
        logger.info(f"[RESUME_PARSER] DOCX parsing extracted {len(text)} characters")
        
        if not text or len(text.strip()) < 10:
            logger.warning(f"[RESUME_PARSER] DOCX file appears empty - only {len(text.strip())} characters extracted")
            raise ValueError("DOCX file appears to be empty or contains no extractable text. Please ensure your document contains text content.")
        
        logger.info(f"[RESUME_PARSER] DOCX parsing completed successfully")
        return extract_resume_data(text)
        
    except ImportError:
        raise  # Re-raise ImportError as-is
    except ValueError:
        raise  # Re-raise ValueError as-is
    except Exception as e:
        error_msg = str(e).lower()
        logger.error(f"[RESUME_PARSER] DOCX parsing failed: {str(e)}")
        if "not a docx" in error_msg or "invalid" in error_msg or "corrupt" in error_msg:
            raise ValueError("The file is not a valid DOCX document")
        raise Exception(f"Error parsing DOCX: {str(e)}")


def extract_resume_data(text: str) -> Dict[str, Any]:
    """
    Extract structured data from resume text
    Returns: {name, email, skills, experience}
    """
    text_lower = text.lower()
    
    # Extract name (usually in first few lines)
    name = extract_name(text)
    
    # Extract email
    email = extract_email(text)
    
    # Extract skills
    skills = extract_skills(text, text_lower)
    
    # Extract experience
    experience = extract_experience(text, text_lower)
    
    return {
        "name": name,
        "email": email,
        "skills": skills,
        "experience": experience,
        "text_length": len(text)
    }


def extract_name(text: str) -> Optional[str]:
    """Extract name from resume (usually at the top)"""
    lines = text.split('\n')[:15]  # Check first 15 lines
    for line in lines:
        line = line.strip()
        if 3 < len(line) < 50:  # Reasonable name length
            # Check if it looks like a name
            if re.match(r'^[A-Za-z\s\.\-]+$', line):
                # Exclude common non-name patterns
                excluded = ['email', 'phone', 'address', 'resume', 'cv', 'linkedin', 'github', 'objective', 'summary']
                if not any(word.lower() in excluded for word in line.split()):
                    if '@' not in line and not re.match(r'^[\d\s\-\+\(\)]+$', line):
                        return line.title()
    return None


def extract_email(text: str) -> Optional[str]:
    """Extract email address using regex"""
    email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
    matches = re.findall(email_pattern, text)
    if matches:
        return matches[0].lower()
    return None


def extract_skills(text: str, text_lower: str) -> List[str]:
    """Extract skills from resume text"""
    skill_keywords = [
        # Programming Languages
        'python', 'java', 'javascript', 'typescript', 'c++', 'c#', 'go', 'rust', 'php', 'ruby',
        'swift', 'kotlin', 'scala', 'r', 'matlab', 'sql', 'html', 'css', 'sass', 'less',
        # Frameworks & Libraries
        'react', 'angular', 'vue', 'node.js', 'express', 'django', 'flask', 'fastapi',
        'spring', 'laravel', 'rails', 'tensorflow', 'pytorch', 'pandas', 'numpy',
        # Tools & Technologies
        'docker', 'kubernetes', 'aws', 'azure', 'gcp', 'git', 'jenkins', 'ci/cd',
        'mongodb', 'postgresql', 'mysql', 'redis', 'elasticsearch', 'kafka',
        # Other
        'agile', 'scrum', 'devops', 'microservices', 'rest api', 'graphql', 'supabase'
    ]
    
    found_skills = []
    for skill in skill_keywords:
        pattern = rf'\b{re.escape(skill)}\b'
        if re.search(pattern, text_lower, re.IGNORECASE):
            # Format skill name properly
            if '.' in skill:
                skill_formatted = skill.upper()
            elif skill == 'ci/cd':
                skill_formatted = 'CI/CD'
            else:
                skill_formatted = skill.title()
            
            if skill_formatted not in found_skills:
                found_skills.append(skill_formatted)
    
    return found_skills[:20]  # Limit to top 20


def extract_experience(text: str, text_lower: str) -> str:
    """
    Extract experience level or summary
    ONLY counts actual work experience, NOT projects, internships, or academic work.
    """
    # First, check for explicit fresher indicators
    fresher_pattern = r'\b(fresher|fresh\s*graduate|no\s*experience|entry\s*level|recent\s*graduate|new\s*graduate)\b'
    if re.search(fresher_pattern, text_lower):
        return "Fresher"
    
    # Look for work experience section headers
    work_experience_section_keywords = [
        r'\b(work\s*experience|professional\s*experience|employment\s*history|work\s*history|career\s*history|experience\s*section)\b',
        r'\b(experience|employment|work\s*history)\s*:',
    ]
    
    # Check if there's a work experience section
    has_work_experience_section = False
    for pattern in work_experience_section_keywords:
        if re.search(pattern, text_lower):
            has_work_experience_section = True
            break
    
    # Check for company/role patterns indicating actual employment
    employment_indicators = [
        r'\b(company|employer|organization|corporation|firm)\s*:',
        r'\b(worked\s*at|employed\s*at|position\s*at|role\s*at|job\s*at)\b',
        r'\b(software\s*engineer|developer|analyst|manager|engineer|consultant)\s*(?:at|in|with)\b',
    ]
    
    has_employment_indicators = False
    for pattern in employment_indicators:
        if re.search(pattern, text_lower):
            has_employment_indicators = True
            break
    
    # If no work experience section or employment indicators found, return Fresher
    if not has_work_experience_section and not has_employment_indicators:
        return "Fresher"
    
    # Patterns for years of experience ONLY in work/professional context
    work_experience_patterns = [
        r'\b(\d+)\s*(?:years?|yrs?|y\.?)\s*(?:of\s*)?(?:work|professional|industry|relevant)\s*experience',
        r'(?:work|professional|industry|relevant)\s*experience[:\s]+(\d+)\s*(?:years?|yrs?)',
        r'\b(\d+)\s*(?:years?|yrs?)\s*(?:of\s*)?experience\s*(?:in|with|at)\s*(?:software|development|engineering|technology)',
    ]
    
    max_years = 0
    for pattern in work_experience_patterns:
        matches = re.finditer(pattern, text_lower, re.IGNORECASE)
        for match in matches:
            try:
                years = int(match.group(1))
                max_years = max(max_years, years)
            except (IndexError, ValueError):
                continue
    
    if max_years > 0:
        return f"{max_years} years"
    
    # If work experience section exists but no years found, check for job titles
    # BUT only if we're in a work experience section context
    if has_work_experience_section or has_employment_indicators:
        # Look for senior/lead roles in work context
        senior_pattern = r'\b(senior|lead|principal|architect|manager|director)\s+(?:software|engineer|developer|analyst|consultant)'
        if re.search(senior_pattern, text_lower):
            return "5+ years"
    
    # Default: No valid work experience found
    return "Fresher"

